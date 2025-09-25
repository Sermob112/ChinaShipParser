# -*- coding: utf-8 -*-
from __future__ import annotations
import json
import re
from pathlib import Path
from datetime import datetime
from typing import List, Dict, Any, Optional

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_UNDERLINE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def load_json_any(path: Path) -> List[Dict[str, Any]]:
    """
    Поддерживает:
      - обычный JSON-массив
      - один объект (оборачиваем в массив)
      - NDJSON (по объекту в строке)
    """
    raw = path.read_text(encoding="utf-8").strip()
    if not raw:
        return []
    try:
        data = json.loads(raw)
        if isinstance(data, list):
            return [x for x in data if isinstance(x, dict)]
        if isinstance(data, dict):
            return [data]
    except json.JSONDecodeError:
        out = []
        for line in raw.splitlines():
            line = line.strip()
            if not line:
                continue
            try:
                obj = json.loads(line)
                if isinstance(obj, dict):
                    out.append(obj)
            except Exception:
                pass
        return out
    return []


def add_hyperlink(paragraph, url: str, text: Optional[str] = None):
    """
    Добавляет кликабельную ссылку в абзац python-docx.
    """
    text = text or url
    part = paragraph.part
    r_id = part.relate_to(
        url,
        reltype="http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    # <w:hyperlink r:id="rIdX">
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    # <w:r><w:rPr><w:u w:val="single"/><w:color w:val="0000FF"/></w:rPr><w:t>text</w:t></w:r>
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')
    rPr.append(color)

    new_run.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)
    return hyperlink


def norm_text(s: str) -> str:
    s = (s or "").replace("\r", "\n").replace("\xa0", " ")
    s = re.sub(r"[ \t]+", " ", s)
    s = re.sub(r"\s*\n\s*", "\n", s)
    return s.strip()


def is_bullet_line(line: str) -> bool:
    """
    Примитивная эвристика: короткие строки без точки в конце — как пункты списка.
    Также считаем пунктами строки, начинающиеся с типичных ключевых слов.
    """
    l = line.strip(" -–•\t")
    if not l:
        return False
    if len(l) <= 64 and not l.endswith((".", "。", "！", "!", ";", "；")):
        return True
    if re.match(r"^(Ship|Self|Pipe|Cruise|LNG|LPG|Small|Middle|Mini|Jackup|Offshore)\b", l, flags=re.I):
        return True
    return False


class ShipyardsDetailsToDocx:
    def __init__(self, in_json: Path, out_docx: Path, title: str = "Shipyards — Profiles"):
        self.in_json = in_json
        self.out_docx = out_docx
        self.title = title

    def build(self):
        items = load_json_any(self.in_json)
        if not items:
            raise SystemExit(f"Пустой вход: {self.in_json}")

        # сортируем по номеру, если есть
        def _key(it):
            try:
                return int(it.get("no", 10**9))
            except Exception:
                return 10**9
        items_sorted = sorted(items, key=_key)

        doc = Document()
        self._setup_styles(doc)

        # Титульная страница
        h = doc.add_heading(self.title, level=0)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p_meta = doc.add_paragraph()
        p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_meta.add_run(f"Всего верфей: {len(items_sorted)}  •  Сформировано: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        run.font.size = Pt(10)

        doc.add_paragraph()  # пустая строка

        # Простой оглавительный список (не автоген TOC, но помогает навигации)
        doc.add_paragraph("Содержание:", style="Heading 2")
        for it in items_sorted:
            no = it.get("no")
            name = it.get("name") or "(без названия)"
            doc.add_paragraph(f"{no}. {name}", style="List Number")

        doc.add_page_break()

        # Основное содержимое
        for idx, it in enumerate(items_sorted, 1):
            no = it.get("no")
            name = it.get("name") or "(без названия)"
            link = it.get("link") or ""
            details = norm_text(it.get("details") or "")

            # Заголовок раздела (H1)
            heading_text = f"{no}. {name}" if no is not None else name
            doc.add_heading(heading_text, level=1)

            # Ссылка
            if link:
                p_link = doc.add_paragraph()
                p_link.add_run("Ссылка: ")
                add_hyperlink(p_link, link, text=link)

            # Текст «как статья» с распознаванием пунктов
            if details:
                lines = details.split("\n")

                # Группируем последовательности буллетов в один «список»
                in_list = False
                for ln in lines:
                    ln = ln.strip()
                    if not ln:
                        # пустая строка завершает список
                        in_list = False
                        doc.add_paragraph("")  # пустой абзац
                        continue

                    if is_bullet_line(ln):
                        doc.add_paragraph(ln.lstrip("•-– "), style="List Bullet")
                        in_list = True
                    else:
                        in_list = False
                        doc.add_paragraph(ln)

            # Между верфями ставим пустую строку (страницу — по желанию)
            if idx != len(items_sorted):
                doc.add_paragraph("")

        doc.save(self.out_docx)
        print(f"[DOCX] Готово: {self.out_docx}")

    def _setup_styles(self, doc: Document):
        # Базовый шрифт
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)

        # Заголовки чуть покрупнее
        for lvl in range(1, 4):
            sname = f'Heading {lvl}'
            if sname in doc.styles:
                st = doc.styles[sname]
                st.font.name = 'Calibri'
                st.font.size = Pt(13 + (3 - lvl))  # H1=15, H2=14, H3=13

        # Уменьшим отступы у списка
        # (по умолчанию ок, но можно тонко настроить при желании)
        pass


def main():
    base = Path(__file__).resolve().parent
    in_json = base / "shipyards_details.json"
    out_docx = base / "shipyards_details.docx"

    conv = ShipyardsDetailsToDocx(in_json=in_json, out_docx=out_docx,
                                  title="Shipyards — Profiles")
    conv.build()


if __name__ == "__main__":
    main()
